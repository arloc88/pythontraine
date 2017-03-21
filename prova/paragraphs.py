from docx import Document



def main():
    # my code here
   # document = Document()

    #document = Document('1_FR_CVF025.docx')

 #   f = open('1_FR_CVF025.docx', 'rb')
   # document = Document(f)
  #  f.close()
    Text = []
    Text = getText ('prova.docx')
    #for element in Text:
        #print(element)
    print(Text)

def getText(filename):
    doc = Document(filename)
    #fullText = []
    fullText=""
    for para in doc.paragraphs:
        fullText+="PARAGRAPHS: " + para.text + "\n \n"
        #fullText.append(para.text)
    #return '\n'.join(fullText)
    return fullText


if __name__ == "__main__":
    main()