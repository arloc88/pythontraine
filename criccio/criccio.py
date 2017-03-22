import docx
import sys
import codecs
sys.stdout = codecs.getwriter("iso-8859-1")(sys.stdout, 'xmlcharrefreplace')

def parse(file):
    document = docx.Document(file)

    lines = []

    lines.append("<p>")
    for p in document.paragraphs:
        text = p.text
        text = text.replace('\t', ' ')
        words = text.split(" ")
        line = ""

        for s in words:
            token = s.lower()
            #print(token)
            if (token == "if"):
                lines.append(line)
                lines.append("<if/>")
                lines.append("<br/><b style=\"color:red\">if</b><br/>")
                line = ""
            elif (token == "then"):
                lines.append(line)
                lines.append("<then/>")
                lines.append("<br/><b style=\"color:red\">then</b><br/>")
                line = ""
            elif (token == "endif"):
                lines.append(line)
                lines.append("<endif/>")
                lines.append("<br/><b style=\"color:red\">endif</b><br/>")
                line = ""
            elif (token == "else"):
                lines.append(line)
                lines.append("<else/>")
                lines.append("<br/><b style=\"color:red\">else</b><br/>")
                line = ""
            else:
                line += s + " "

        if (line != ""):
            lines.append(line)

        lines.append("</p>")

    out_file = open(file+".html", "w", encoding="utf-8")
    for line in lines:
        out_file.write(line)
        out_file.write("\n")
    out_file.close()

    '''
    def parse(file):
    document = docx.Document(file)

    #lines = []
    tree = Tree()
    root = tree.create_node("root",0)
    node = root

    i = 1
    for p in document.paragraphs:
        text = p.text
        words = text.split(" ")

        line = ""
        for s in words:
            token = s.lower()
            #print(token)
            if (token == "if" or token == "when"):
                #printSpace ("<if>");
                tree.create_node(line, i, parent=node.identifier)
                line = ""
                i+=1
                node = tree.create_node("if", i, parent=root.identifier)
            elif (token == "then"):
                tree.create_node(line, i, parent=node.identifier)
                line = ""
                i += 1
                node = tree.create_node("then", i, parent=node.identifier)
            elif (token == "endif"):
                tree.create_node(line, i, parent=node.identifier)
                line = ""
                i += 1
                tree.create_node("endif", i, parent=node.identifier)
            elif (token == "else"):
                #printSpace("<else>")
                tree.create_node(line, i, parent=node.identifier)
                line = ""
                i += 1
                tree.create_node("else", i, parent=node.identifier)
            else:
                line += token + " "

            i += 1

    tree.show()

    '''

def main():
    parse("prova.docx")
    
if __name__ == '__main__':
    main()
